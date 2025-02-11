import pytest
from fastapi.testclient import TestClient
from app.main import app
import docx
from io import BytesIO
from unittest.mock import patch, Mock
from datetime import datetime
from app.schemas.contract import ContractAnalysis, Issue, Location, Suggestion

@pytest.fixture
def client():
    """Create a test client."""
    return TestClient(app)

@pytest.fixture
def sample_docx():
    """Create a sample DOCX file for testing."""
    doc = docx.Document()
    doc.add_paragraph("Liability Clause")
    doc.add_paragraph("The contractor shall be liable for all damages.")
    doc.add_paragraph("Payment Terms")
    doc.add_paragraph("Payment shall be made within 30 days.")
    
    # Save to BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.read()

def test_health_check(client):
    """Test the health check endpoint."""
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "healthy"}

@patch('app.services.claude_service.ClaudeService.analyze_contract')
def test_analyze_contract_success(mock_analyze, client, sample_docx):
    """Test successful contract analysis."""
    # Create a ContractAnalysis instance for the mock
    analysis = ContractAnalysis(
        issues=[
            Issue(
                type="liability_clause",
                severity="high",
                description="Unlimited liability clause detected",
                location=Location(
                    paragraph=2,
                    text="The contractor shall be liable for all damages."
                ),
                suggestion="Consider adding liability cap"
            )
        ],
        suggestions=[
            Suggestion(
                category="payment_terms",
                description="Payment terms could be more specific",
                current="Payment shall be made within 30 days.",
                suggested="Payment shall be made within 30 days of invoice receipt"
            )
        ],
        risk_score=75,
        analysis_timestamp=datetime.utcnow()
    )
    
    # Set up the mock to return the ContractAnalysis instance
    mock_analyze.return_value = analysis
    
    # Make request
    response = client.post(
        "/api/v1/analyze-contract",
        files={"file": ("test.docx", sample_docx)},
        data={"jurisdiction": "US", "contract_type": "service"}
    )
    
    # Verify response
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "success"
    assert "analysis" in data
    assert len(data["analysis"]["issues"]) == 1
    assert len(data["analysis"]["suggestions"]) == 1
    assert data["analysis"]["risk_score"] == 75

def test_analyze_contract_invalid_file(client):
    """Test analyzing an invalid file type."""
    response = client.post(
        "/api/v1/analyze-contract",
        files={"file": ("test.txt", b"Not a DOCX file")},
    )
    
    assert response.status_code == 400
    assert "Only DOCX files are supported" in response.json()["detail"]

def test_analyze_contract_file_too_large(client):
    """Test analyzing a file that exceeds size limit."""
    # Create a large file
    large_file = b"x" * (10 * 1024 * 1024 + 1)  # 10MB + 1 byte
    
    response = client.post(
        "/api/v1/analyze-contract",
        files={"file": ("test.docx", large_file)},
    )
    
    assert response.status_code == 400
    assert "File size exceeds maximum limit" in response.json()["detail"]

@patch('app.services.claude_service.ClaudeService.analyze_contract')
def test_analyze_contract_claude_error(mock_analyze, client, sample_docx):
    """Test handling of Claude API errors."""
    # Mock Claude service to raise an exception
    mock_analyze.side_effect = Exception("Claude API Error")
    
    response = client.post(
        "/api/v1/analyze-contract",
        files={"file": ("test.docx", sample_docx)},
    )
    
    assert response.status_code == 200  # We return 200 with error in response body
    data = response.json()
    assert data["status"] == "error"
    assert data["error"]["error_code"] == "PROCESSING_ERROR" 
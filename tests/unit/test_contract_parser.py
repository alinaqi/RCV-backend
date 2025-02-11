import pytest
from fastapi import UploadFile, HTTPException
from app.services.contract_parser import ContractParser
from io import BytesIO
import docx

@pytest.fixture
def sample_docx():
    """Create a sample DOCX file for testing."""
    doc = docx.Document()
    doc.add_paragraph("Liability Clause")
    doc.add_paragraph("The contractor shall be liable for all damages.")
    doc.add_paragraph("Payment Terms")
    doc.add_paragraph("Payment shall be made within 30 days.")
    doc.add_paragraph("Notice Period")
    doc.add_paragraph("A notice period of 30 days is required.")
    
    # Save to BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

@pytest.fixture
def upload_file(sample_docx):
    """Create a FastAPI UploadFile from the sample DOCX."""
    return UploadFile(
        filename="test.docx",
        file=sample_docx
    )

@pytest.mark.asyncio
async def test_parse_docx_valid_file(upload_file):
    """Test parsing a valid DOCX file."""
    parser = ContractParser()
    result = await parser.parse_docx(upload_file)
    
    assert "[P1] Liability Clause" in result
    assert "[P2] The contractor shall be liable for all damages." in result
    assert "[P3] Payment Terms" in result
    assert "[P4] Payment shall be made within 30 days." in result

@pytest.mark.asyncio
async def test_parse_docx_invalid_file():
    """Test parsing an invalid file raises HTTPException."""
    invalid_file = UploadFile(
        filename="test.txt",
        file=BytesIO(b"Not a DOCX file")
    )
    
    parser = ContractParser()
    with pytest.raises(HTTPException) as exc_info:
        await parser.parse_docx(invalid_file)
    
    assert exc_info.value.status_code == 400
    assert "Failed to parse DOCX file" in str(exc_info.value.detail)

def test_extract_sections():
    """Test extracting sections from parsed text."""
    parser = ContractParser()
    text = """
    [P1] Liability Clause
    [P2] The contractor shall be liable for all damages.
    [P3] Payment Terms
    [P4] Payment shall be made within 30 days.
    [P5] Notice Period
    [P6] A notice period of 30 days is required.
    """
    
    sections = parser.extract_sections(text)
    
    assert "liability" in sections["liability_clauses"].lower()
    assert "payment" in sections["payment_terms"].lower()
    assert "notice" in sections["notice_periods"].lower()

def test_extract_sections_empty_text():
    """Test extracting sections from empty text."""
    parser = ContractParser()
    sections = parser.extract_sections("")
    
    assert all(value == "" for value in sections.values()) 
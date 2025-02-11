from typing import Dict, Optional, List, Tuple
import docx
from fastapi import UploadFile, HTTPException
import logging
from io import BytesIO
import re
import anthropic
from src.app.core.config import settings
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

class RedlineItem:
    """Class to represent a redline (tracked change) in a document."""
    def __init__(self, paragraph_number: int, original_text: str, modified_text: str, author: str, date: str, change_type: str):
        self.paragraph_number = paragraph_number
        self.original_text = original_text
        self.modified_text = modified_text
        self.author = author
        self.date = date
        self.change_type = change_type  # 'insertion', 'deletion', or 'modification'

class ParsedDocument:
    """Class to represent a parsed document with its text and redlines."""
    def __init__(self, text: str, redlines: List[RedlineItem]):
        self.text = text
        self.redlines = redlines

class ContractParser:
    """Service for parsing DOCX contract files."""
    
    def __init__(self):
        self.claude_client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
    
    # Common contract keywords and phrases that indicate a legal contract
    CONTRACT_INDICATORS = [
        r"agreement",
        r"contract",
        r"terms and conditions",
        r"parties",
        r"hereby agree",
        r"obligations",
        r"effective date",
        r"in witness whereof",
        r"signature",
        r"signed by",
    ]
    
    # Essential sections that should be present in a valid contract
    ESSENTIAL_SECTIONS = [
        r"parties?",
        r"purpose|scope",
        r"terms?",
        r"conditions?",
        r"obligations?",
        r"payment|compensation",
        r"termination",
        r"governing law|jurisdiction",
        r"signature|execution",
    ]
    
    def _extract_redlines(self, paragraph: Paragraph, paragraph_number: int) -> List[RedlineItem]:
        """
        Extract redline items from a paragraph.
        
        Args:
            paragraph (Paragraph): The paragraph to analyze
            paragraph_number (int): The number of the paragraph
            
        Returns:
            List[RedlineItem]: List of redline items found in the paragraph
        """
        redlines = []
        
        if not hasattr(paragraph._element, 'xpath'):
            return redlines

        # Define namespace map
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        try:
            # Find deletions
            dels = paragraph._element.findall('.//w:del', nsmap)
            for deletion in dels:
                try:
                    author = deletion.get(f'{{{nsmap["w"]}}}author', 'Unknown')
                    date = deletion.get(f'{{{nsmap["w"]}}}date', 'Unknown')
                    text = deletion.text if deletion.text else ""
                    
                    if text:
                        redline = RedlineItem(
                            paragraph_number=paragraph_number,
                            original_text=text,
                            modified_text="",
                            author=author,
                            date=date,
                            change_type='deletion'
                        )
                        redlines.append(redline)
                except Exception as e:
                    logger.error(f"Error processing deletion in paragraph {paragraph_number}: {str(e)}", exc_info=True)
                    continue

            # Find insertions
            ins = paragraph._element.findall('.//w:ins', nsmap)
            for insertion in ins:
                try:
                    author = insertion.get(f'{{{nsmap["w"]}}}author', 'Unknown')
                    date = insertion.get(f'{{{nsmap["w"]}}}date', 'Unknown')
                    text = insertion.text if insertion.text else ""
                    
                    if text:
                        redline = RedlineItem(
                            paragraph_number=paragraph_number,
                            original_text="",
                            modified_text=text,
                            author=author,
                            date=date,
                            change_type='insertion'
                        )
                        redlines.append(redline)
                except Exception as e:
                    logger.error(f"Error processing insertion in paragraph {paragraph_number}: {str(e)}", exc_info=True)
                    continue

            # Find moves (treated as modifications)
            moves_from = paragraph._element.findall('.//w:moveFrom', nsmap)
            moves_to = paragraph._element.findall('.//w:moveTo', nsmap)
            
            for move in moves_from + moves_to:
                try:
                    author = move.get(f'{{{nsmap["w"]}}}author', 'Unknown')
                    date = move.get(f'{{{nsmap["w"]}}}date', 'Unknown')
                    text = move.text if move.text else ""
                    
                    if text:
                        change_type = 'deletion' if move.tag.endswith('moveFrom') else 'insertion'
                        redline = RedlineItem(
                            paragraph_number=paragraph_number,
                            original_text=text if change_type == 'deletion' else "",
                            modified_text=text if change_type == 'insertion' else "",
                            author=author,
                            date=date,
                            change_type=change_type
                        )
                        redlines.append(redline)
                except Exception as e:
                    logger.error(f"Error processing move in paragraph {paragraph_number}: {str(e)}", exc_info=True)
                    continue
                    
        except Exception as e:
            logger.error(f"Error extracting redlines from paragraph {paragraph_number}: {str(e)}", exc_info=True)
            
        return redlines

    @staticmethod
    async def parse_docx(file: UploadFile) -> ParsedDocument:
        """
        Parse DOCX file into plain text while maintaining structure and extracting redlines.
        
        Args:
            file (UploadFile): The uploaded DOCX file
            
        Returns:
            ParsedDocument: Object containing formatted text and redline items
            
        Raises:
            HTTPException: If file parsing fails
        """
        try:
            content = await file.read()
            doc = docx.Document(BytesIO(content))
            
            # Extract text with paragraph numbers and redlines
            text_blocks = []
            all_redlines = []
            parser = ContractParser()  # Create instance to access _extract_redlines
            
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    text_blocks.append(f"[P{i}] {paragraph.text}")
                    # Extract redlines from this paragraph
                    redlines = parser._extract_redlines(paragraph, i)
                    all_redlines.extend(redlines)
            
            return ParsedDocument(
                text="\n\n".join(text_blocks),
                redlines=all_redlines
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=400,
                detail="Failed to parse DOCX file. Please ensure it's a valid DOCX document."
            )
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        """
        Extract key contract sections.
        
        Args:
            text (str): The parsed contract text
            
        Returns:
            Dict[str, str]: Dictionary of section name to section content
        """
        sections = {
            "liability_clauses": "",
            "payment_terms": "",
            "notice_periods": "",
            "termination_clauses": "",
            "governing_law": ""
        }
        
        current_section = None
        section_content = []
        
        # Simple keyword-based section detection
        keywords = {
            "liability": "liability_clauses",
            "indemnif": "liability_clauses",
            "payment": "payment_terms",
            "fee": "payment_terms",
            "notice": "notice_periods",
            "notif": "notice_periods",
            "terminat": "termination_clauses",
            "govern": "governing_law",
            "law": "governing_law",
            "jurisdiction": "governing_law"
        }
        
        for line in text.split("\n"):
            line = line.lower()
            
            # Check if line contains any section keywords
            for keyword, section in keywords.items():
                if keyword in line:
                    current_section = section
                    section_content = []
                    break
            
            if current_section and line.strip():
                section_content.append(line)
                sections[current_section] = "\n".join(section_content)
        
        return sections
    
    async def _validate_with_claude(self, text: str) -> bool:
        """
        Use Claude to validate if the document is a legitimate legal contract.
        
        Args:
            text (str): The parsed contract text
            
        Returns:
            bool: True if Claude determines this is a valid contract
        """
        try:
            prompt = """You are a legal document validator. Analyze the following document and determine if it is a legitimate legal contract. 
            Consider these aspects:
            1. Does it have proper contract structure and formatting?
            2. Does it contain essential contract elements?
            3. Is the language formal and legally binding?
            4. Are there clear obligations and agreements?

            Note that user may upload a template contract, which would be a valid contract.
            
            Respond with ONLY 'true' if it's a valid contract, or 'false' if it's not.
            
            Document text:
            {text}"""
            
            message = self.claude_client.messages.create(
                model=settings.CLAUDE_MODEL,
                max_tokens=settings.MAX_TOKENS,
                temperature=0,
                system="You are a legal document validator. Respond with ONLY 'true' or 'false'.",
                messages=[
                    {
                        "role": "user",
                        "content": prompt.format(text=text)
                    }
                ]
            )
            
            response = message.content[0].text.strip().lower()
            return response == 'true'
            
        except Exception as e:
            logger.error(f"Error validating contract with Claude: {str(e)}", exc_info=True)
            # Fall back to basic validation if Claude fails
            return self._basic_validation(text)
    
    def _basic_validation(self, text: str) -> bool:
        """
        Basic validation using regex patterns as a fallback.
        
        Args:
            text (str): The parsed contract text
            
        Returns:
            bool: True if the document appears to be a valid contract
        """
        text_lower = text.lower()
        
        # Check for contract indicators
        indicator_count = sum(1 for indicator in self.CONTRACT_INDICATORS 
                            if re.search(indicator, text_lower))
        
        # Check for essential sections
        section_count = sum(1 for section in self.ESSENTIAL_SECTIONS 
                          if re.search(section, text_lower))
        
        return indicator_count >= 3 and section_count >= 4
    
    async def is_valid_contract(self, text: str) -> bool:
        """
        Validate if the document appears to be a legitimate contract using Claude AI.
        Falls back to basic validation if Claude is unavailable.
        
        Args:
            text (str): The parsed contract text
            
        Returns:
            bool: True if the document appears to be a valid contract
        """
        try:
            return await self._validate_with_claude(text)
        except Exception as e:
            logger.error(f"Failed to validate with Claude, falling back to basic validation: {str(e)}", exc_info=True)
            return self._basic_validation(text)